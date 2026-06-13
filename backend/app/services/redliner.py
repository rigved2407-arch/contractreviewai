import difflib
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from lxml import etree

RED = RGBColor(0xFF, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x80, 0x00)
MATCH_THRESHOLD = 0.6


def _normalize(text: str) -> str:
    import re
    return re.sub(r'\s+', ' ', text).strip()


def generate_redlined_docx(
    original_path: str,
    clauses: list[dict],
    output_path: str,
) -> str:
    doc = Document(original_path)
    para_texts = [p.text for p in doc.paragraphs]

    for clause in clauses:
        suggested = clause.get("suggested_redline", "")
        orig_text = clause.get("clause_text", "")
        if not suggested or not orig_text:
            continue

        best_match_idx, best_ratio = _find_best_match(orig_text, para_texts)

        if best_match_idx is not None and best_ratio >= MATCH_THRESHOLD:
            _apply_fuzzy_redline(doc, doc.paragraphs[best_match_idx], suggested, orig_text)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def _find_best_match(needle: str, haystacks: list[Optional[str]]) -> tuple[Optional[int], float]:
    needle_norm = _normalize(needle)
    best_idx = None
    best_ratio = 0.0
    for i, h in enumerate(haystacks):
        if not h:
            continue
        ratio = difflib.SequenceMatcher(None, needle_norm, _normalize(h)).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_idx = i
    return best_idx, best_ratio


def _apply_fuzzy_redline(doc: Document, para, suggested: str, orig_text: str):
    if doc.paragraphs.index(para) < len(doc.paragraphs) - 1:
        next_para = doc.paragraphs[doc.paragraphs.index(para) + 1]
        if "[SUGGESTED]" in (next_para.text or ""):
            return

    orig_short = orig_text[:100] if len(orig_text) > 100 else orig_text
    for run in para.runs:
        run.font.color.rgb = RED
        run.font.strike = True

    new_para = doc.add_paragraph()
    new_para.paragraph_format.space_before = Pt(2)
    new_para.paragraph_format.space_after = Pt(2)
    new_run = new_para.add_run(f"[SUGGESTED: {suggested}]")
    new_run.font.color.rgb = GREEN
    new_run.font.size = Pt(10)
    new_run.font.italic = True

    comment_run = para.add_run(f"\n[REDLINED: {orig_short}]")
    comment_run.font.color.rgb = RED
    comment_run.font.size = Pt(9)
    comment_run.font.italic = True
