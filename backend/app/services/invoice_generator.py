from datetime import datetime
from pathlib import Path

from app.config import settings


def generate_gst_invoice(
    invoice_number: str,
    org_name: str,
    org_gstin: str,
    org_address: str,
    org_email: str,
    plan_name: str,
    amount_inr: int,
    hsn_code: str = "998313",
) -> str:
    """Generate a GST-compliant invoice PDF and return the file path.

    HSN 998313 = Legal advisory and representation services.
    CGST 9% + SGST 9% = 18% GST.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    base_amount = round(amount_inr / 1.18)
    cgst = round(base_amount * 0.09)
    sgst = round(base_amount * 0.09)
    total = base_amount + cgst + sgst

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TAX INVOICE")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell_left = table.cell(0, 0)
    cell_right = table.cell(0, 1)

    cell_left.text = f"""Seller:
Contract Review AI
GSTIN: {org_gstin or 'Not provided'}
Address: Registered Office, India
Email: billing@contractreviewai.com"""

    cell_right.text = f"""Buyer:
{org_name}
GSTIN: {org_gstin or 'Not provided'}
Address: {org_address or 'N/A'}
Email: {org_email}"""

    doc.add_paragraph()

    doc.add_paragraph(f"Invoice No: {invoice_number}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%b-%Y')}")
    doc.add_paragraph(f"Place of Supply: Maharashtra (27)")

    doc.add_paragraph()
    p = doc.add_paragraph("Description of Services")
    p.runs[0].bold = True

    items_table = doc.add_table(rows=2, cols=6)
    items_table.style = "Light Grid Accent 1"

    headers = ["#", "HSN/SAC", "Description", "Amount (₹)", "CGST 9%", "SGST 9%"]
    for i, h in enumerate(headers):
        items_table.cell(0, i).text = h

    items_table.cell(1, 0).text = "1"
    items_table.cell(1, 1).text = hsn_code
    items_table.cell(1, 2).text = f"Contract Review AI - {plan_name} Plan (Monthly Subscription)"
    items_table.cell(1, 3).text = f"{base_amount:,.0f}"
    items_table.cell(1, 4).text = f"{cgst:,.0f}"
    items_table.cell(1, 5).text = f"{sgst:,.0f}"

    doc.add_paragraph()
    doc.add_paragraph(f"Total Amount (excl. GST): ₹{base_amount:,.0f}")
    doc.add_paragraph(f"CGST @ 9%: ₹{cgst:,.0f}")
    doc.add_paragraph(f"SGST @ 9%: ₹{sgst:,.0f}")
    p = doc.add_paragraph(f"Total Invoice Value: ₹{total:,.0f}")
    p.runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph("Amount in words: Rupees " + _number_to_words(total) + " only")
    doc.add_paragraph()
    doc.add_paragraph(
        "Declaration: This is a system-generated GST-compliant invoice. "
        "Valid under CGST/SGST Act, 2017. E-invoice applicable for turnover > INR 5 Cr."
    )

    out_dir = Path(settings.storage_dir) / "invoices"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = str(out_dir / f"{invoice_number}.docx")
    doc.save(out_path)
    return out_path


def _number_to_words(n: int) -> str:
    if n == 0:
        return "Zero"
    ones = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine"]
    tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
    teens = ["Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen",
             "Seventeen", "Eighteen", "Nineteen"]

    def _convert(num):
        if num < 10:
            return ones[num]
        if num < 20:
            return teens[num - 10]
        if num < 100:
            return tens[num // 10] + (" " + ones[num % 10] if num % 10 else "")
        if num < 1000:
            return ones[num // 100] + " Hundred" + (" " + _convert(num % 100) if num % 100 else "")
        if num < 100000:
            return _convert(num // 1000) + " Thousand" + (" " + _convert(num % 1000) if num % 1000 else "")
        if num < 10000000:
            return _convert(num // 100000) + " Lakh" + (" " + _convert(num % 100000) if num % 100000 else "")
        return _convert(num // 10000000) + " Crore" + (" " + _convert(num % 10000000) if num % 10000000 else "")

    return _convert(n)
